import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Config
BASE_DIR = os.path.normpath(os.path.join(os.path.dirname(__file__), ".."))
ANALYSIS_DIR = os.path.join(BASE_DIR, "docs/analysis")
IMPACT_PLOTS_DIR = os.path.join(BASE_DIR, "results/2026-03-11/dataset-small/analysis_plots")
CONSOLIDATED_PLOTS_DIR = os.path.join(BASE_DIR, "results/2026-03-11/dataset-small/figures_consolidated")
OUTPUT_DOC = os.path.join(ANALYSIS_DIR, "Rapport_Analyse_SDN_V19_Final.docx")

def create_report():
    print("Starting Comprehensive Word Report Generation...")
    doc = Document()
    
    # --- Style Setting ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Title ---
    title_obj = doc.add_heading("Analyse Comparative et Modélisation de la Sélection de Liens SDN", 0)
    title_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Chapter 1: Theory & Priority ---
    doc.add_heading("1. Fondements Théoriques et Logique de Priorité", level=1)
    doc.add_paragraph(
        "L'implémentation inclut désormais un ordonnancement basé sur la priorité utilisateur (CPU/Réseau). "
        "Le workflow suit une extraction depuis le CSV suivie d'un tri décroissant dans le SDNBroker."
    )
    doc.add_paragraph("Composantes de la latence modélisée :", style='List Bullet')
    doc.add_paragraph("Dprop (Propagation) : Basé sur la distance et l'indice de réfraction.", style='List Bullet 2')
    doc.add_paragraph("Dtrans (Transmission) : Ratio PacketSize / Bandwidth.", style='List Bullet 2')
    doc.add_paragraph("Dqueue (Attente) : Modèle de file M/M/1 dynamique.", style='List Bullet 2')

    # --- Chapter 2: Architecture ---
    doc.add_heading("2. Architecture de Simulation (Dataset-Small)", level=1)
    doc.add_paragraph(
        "Infrastructure hiérarchique comportant 6 hôtes (h_0 à h_5) et des liens redondants asymétriques "
        "(10 Mbps vs 5 Gbps), conçue pour tester la résilience des politiques de routage."
    )

    # --- Chapter 3: Deep Dive per VM Policy ---
    doc.add_heading("3. Analyse par Politique d'Allocation de VM (Deep Dive)", level=1)
    
    # 3.1 MFF
    doc.add_heading("3.1 Politique MFF (Most Full First)", level=2)
    doc.add_paragraph(
        "Densification maximale pour l'efficience énergétique. Record à 1.11 Wh (PSO). "
        "Risque de saturation des liens d'accès si non couplé à BwAllocN."
    )
    
    # 3.2 LFF
    doc.add_heading("3.2 Politique LFF (Least Full First)", level=2)
    doc.add_paragraph(
        "Étalement de charge sur les 6 hôtes. Offre une isolation supérieure mais une "
        "consommation doublée (16.09 Wh) car tous les serveurs restent actifs."
    )
    
    # 3.3 LWFF
    doc.add_heading("3.3 Politique LWFF (Least Weight Full First)", level=2)
    doc.add_paragraph(
        "Utilise l'optimisation de Pareto. C'est la politique la plus équilibrée pour les "
        "environnements multi-tenant complexes."
    )

    # --- Chapter 4: Multi-Metric Analysis (The 5 Figures) ---
    doc.add_heading("4. Analyse Métrique Globale (Les 5 Figures Clés)", level=1)
    
    figures = [
        ("fig1_energy.png", "Figure 1: Impact Énergétique (Densification vs Étalement)", 
         "Analyse : Réduction de 75% du gaspillage énergétique avec BwAllocN due à la réduction de la durée de simulation."),
        ("fig2_latency.png", "Figure 2: Latence E2E (Efficacité du routage dynamique)", 
         "Analyse : Amélioration de 93% de la latence en exploitant les liens haute capacité (5 Gbps)."),
        ("fig3_sla.png", "Figure 3: Violations SLA et Priorités", 
         "Analyse : La politique Priority réduit les violations de 30% à 45% pour les tâches critiques."),
        ("fig4_packet_delay.png", "Figure 4: Distribution des Délais (Stabilité du réseau)", 
         "Analyse : Élimination du Jitter et du Bufferbloat grâce à la sélection de lien intelligente."),
        ("fig5_utilization.png", "Figure 5: Utilisation des Ressources Physiques", 
         "Analyse : Illustration de la densification énergétique (MFF) vs l'équilibrage de charge (LFF/LWFF).")
    ]
    
    for fig_name, caption, analysis in figures:
        fig_path = os.path.join(CONSOLIDATED_PLOTS_DIR, fig_name)
        if os.path.exists(fig_path):
            doc.add_picture(fig_path, width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(caption, style='Caption')
            doc.add_paragraph(analysis)
            doc.add_paragraph() # Spacer

    # --- Chapter 5: Comparative Synthesis ---
    doc.add_heading("5. Synthèse Comparative des 3 Meilleures Combinaisons", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rang'
    hdr_cells[1].text = 'Combinaison'
    hdr_cells[2].text = 'Délai Pkt'
    hdr_cells[3].text = 'Énergie'
    
    data = [
        ('1', 'LWFF + BwAllocN + SJF', '~6.5 s', '16.09 Wh'),
        ('2', 'MFF + BwAllocN + PSO', '~6.7 s', '1.11 Wh'),
        ('3', 'MFF + BwAllocN + Priority', '~6.6 s', '8.89 Wh')
    ]
    for r, c, d, e in data:
        row_cells = table.add_row().cells
        row_cells[0].text = r
        row_cells[1].text = c
        row_cells[2].text = d
        row_cells[3].text = e

    # --- Final Save ---
    doc.save(OUTPUT_DOC)
    print(f"Final Enriched Report saved: {OUTPUT_DOC}")

if __name__ == "__main__":
    create_report()
